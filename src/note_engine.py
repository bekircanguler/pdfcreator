#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Kurumsal Bilgi Notu — Word (.docx) + PDF üretim motoru
"""

import io
import sys
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional

from engine import (
    get_fonts, get_title_fonts, get_italic_fonts, hex_rgb, wrap_text,
    PAGE_W, PAGE_H,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from PIL import Image, ImageOps
from pypdf import PdfWriter, PdfReader

# ─── Sabitler ──────────────────────────────────────────────────────────────────
BIRIM_ADI = "Fen İşleri Dairesi Başkanlığı\nBakım ve Onarım Şube Müdürlüğü"
BELGE_ADI = "BİLGİ NOTU"
LOGO_FILE = "logo.png"

MARGIN      = 20 * mm
CONTENT_W   = PAGE_W - 2 * MARGIN
LINE_H_BODY = 15      # pt — body satır yüksekliği
LINE_H_SM   = 11      # pt — küçük satır

BODY_SIZE   = 11      # pt — gövde metni (Arial Unicode, 11pt → temiz A4 görünümü)
HEAD_SIZE   = 11      # pt — bölüm başlıkları (bold)

# Dinamik saha notu
_SAHA_KEYWORDS = [
    "boya", "tadilat", "inşaat", "sıva", "alçı", "zemin", "çatı", "duvar",
]
_SAHA_NOTE = (
    "Not: Bahse konu inşaat işleri kapsamında yer alan boya/tadilat işleri "
    "kendi marifetimizle yürütülecek olup, diğer imalatlar hizmet alımı "
    "yoluyla karşılanacaktır."
)

IMAGE_EXTS = frozenset({".jpg", ".jpeg", ".png"})


# ──────────────────────────────────────────────────────────────────────────────
# Yardımcılar
# ──────────────────────────────────────────────────────────────────────────────

def _logo_path() -> Optional[Path]:
    if hasattr(sys, "_MEIPASS"):
        p = Path(sys._MEIPASS) / LOGO_FILE
    else:
        p = Path(__file__).resolve().parent.parent / LOGO_FILE
    return p if p.exists() else None


def _today_str() -> str:
    d = date.today()
    return d.strftime("%d.%m.%Y")


def _safe_fname(konu: str) -> str:
    invalid = r'\/:*?"<>|'
    name = "".join(c if c not in invalid else "_" for c in konu.replace(" ", "_"))
    tarih = date.today().strftime("%Y%m%d")
    return (name.strip(". ") or "Bilgi_Notu") + f"_{tarih}"


def _need_saha_note(disciplines: List[Dict]) -> bool:
    for disc in disciplines:
        alt = disc.get("alt_isi", "").lower()
        if any(kw in alt for kw in _SAHA_KEYWORDS):
            return True
    return False


def _scan_photos(folder: str) -> List[Path]:
    return sorted(
        f for f in Path(folder).iterdir()
        if f.is_file() and f.suffix.lower() in IMAGE_EXTS
    )


# ──────────────────────────────────────────────────────────────────────────────
# PDF Motor
# ──────────────────────────────────────────────────────────────────────────────

class _PDFPainter:
    def __init__(self, dest):
        get_fonts()
        get_title_fonts()
        self._c = rl_canvas.Canvas(dest, pagesize=A4)
        self._pages: List[List] = [[]]
        self._page = 0
        self.y = PAGE_H - MARGIN
        # Gövde metni: Arial Unicode — Türkçe karakter desteği garantili
        self._body_reg, self._body_bold = get_fonts()
        self._body_italic, self._body_bold_italic = get_italic_fonts()
        self._title_reg, self._title_bold = get_title_fonts()

    def _new_page(self) -> None:
        self._c.showPage()
        self._pages.append([])
        self._page += 1
        self.y = PAGE_H - MARGIN
        self.draw_mini_header()   # 2. ve sonraki sayfalarda mini başlık

    def _ensure(self, need: float) -> None:
        if self.y - need < MARGIN + 20:
            self._new_page()

    def _hline(self, y: float, width: float = 0.5, color=(0.7, 0.7, 0.7)) -> None:
        c = self._c
        c.setStrokeColorRGB(*color)
        c.setLineWidth(width)
        c.line(MARGIN, y, PAGE_W - MARGIN, y)

    def _rect_fill(self, x, y, w, h, rgb) -> None:
        self._c.setFillColorRGB(*rgb)
        self._c.rect(x, y, w, h, fill=1, stroke=0)

    def draw_main_header(self) -> None:
        c = self._c
        y = PAGE_H - MARGIN

        self._rect_fill(MARGIN, y - 2, CONTENT_W, 2, hex_rgb("#2563EB"))

        logo = _logo_path()
        logo_h = 0.0
        if logo:
            try:
                with Image.open(logo) as im:
                    iw, ih = im.size
                    logo_target_h = 36
                    logo_target_w = int(iw * logo_target_h / ih)
                    reader = ImageReader(str(logo))
                    c.drawImage(reader, MARGIN, y - 2 - logo_target_h,
                                width=logo_target_w, height=logo_target_h,
                                preserveAspectRatio=True, mask="auto")
                    logo_h = logo_target_h + 4
                    text_x = MARGIN + logo_target_w + 8
            except Exception:
                text_x = MARGIN
        else:
            text_x = MARGIN

        c.setFont(self._body_reg, 8)
        c.setFillColorRGB(0.4, 0.4, 0.4)
        for i, line in enumerate(BIRIM_ADI.split("\n")):
            c.drawString(text_x, y - 4 - 10 * (i + 1), line)

        c.setFont(self._title_bold, 14)
        c.setFillColorRGB(*hex_rgb("#1A2E4A"))
        bw = pdfmetrics.stringWidth(BELGE_ADI, self._title_bold, 14)
        c.drawString(PAGE_W - MARGIN - bw, y - 4 - 14, BELGE_ADI)

        sep_y = y - max(logo_h, 32) - 8
        self._hline(sep_y, 1.2, hex_rgb("#2563EB"))
        self.y = sep_y - 10

    def draw_mini_header(self) -> None:
        c = self._c
        y = PAGE_H - MARGIN
        c.setFont(self._body_bold, 8)
        c.setFillColorRGB(0.45, 0.45, 0.45)
        c.drawString(MARGIN, y, BIRIM_ADI.replace("\n", " — ") + " · " + BELGE_ADI)
        self._hline(y - 4, 0.5)
        self.y = y - 14

    def draw_page_number(self, cur: int, total: int) -> None:
        c = self._c
        c.setFont(self._body_reg, 8)
        c.setFillColorRGB(0.55, 0.55, 0.55)
        txt = f"Sayfa {cur} / {total}"
        tw = pdfmetrics.stringWidth(txt, self._body_reg, 8)
        c.drawString(PAGE_W - MARGIN - tw, MARGIN * 0.45, txt)

    def section_title(self, text: str) -> None:
        self._ensure(36)
        c = self._c
        box_h = 20
        box_bottom = self.y - box_h
        # Arka plan kutusu
        self._rect_fill(MARGIN, box_bottom, CONTENT_W, box_h, hex_rgb("#EEF3FF"))
        c.setFont(self._body_bold, HEAD_SIZE)
        c.setFillColorRGB(*hex_rgb("#1A2E4A"))
        # Baseline: kutunun altından 7pt → descender'lar çizgiyle örtüşmez
        c.drawString(MARGIN + 6, box_bottom + 7, text)   # .upper() kaldırıldı
        # Çizgi: kutunun tam alt kenarı
        self._hline(box_bottom, 0.8, hex_rgb("#4A90D9"))
        # İçerik: çizgiden 14pt aşağı → 11pt font ascent (~8pt) + 6pt boşluk
        self.y = box_bottom - 14

    def field_label(self, label: str, value: str = "") -> None:
        self._ensure(LINE_H_BODY)
        c = self._c
        c.setFont(self._body_bold, BODY_SIZE)
        c.setFillColorRGB(*hex_rgb("#374151"))
        c.drawString(MARGIN, self.y, label + ":")
        label_w = pdfmetrics.stringWidth(label + ": ", self._body_bold, BODY_SIZE)
        val_x   = MARGIN + max(label_w, 90)
        val_w   = PAGE_W - MARGIN - val_x
        if not value:
            self.y -= LINE_H_BODY
            return
        c.setFont(self._body_reg, BODY_SIZE)
        c.setFillColorRGB(0.18, 0.18, 0.18)
        lines = wrap_text(value, self._body_reg, BODY_SIZE, val_w)
        for i, ln in enumerate(lines):
            if i > 0:
                self._ensure(LINE_H_BODY)
            c.drawString(val_x, self.y, ln)
            self.y -= LINE_H_BODY

    def multiline_field(self, label: str, text: str) -> None:
        self.field_label(label)
        if text:
            self.body_text(text, indent=8)

    def _justify_line(self, text: str, x: float, y: float,
                       width: float, font: str, size: float) -> None:
        """Tek satırı iki yana yaslar (son satır veya tek kelime ise yaslamaz)."""
        words = text.split()
        if len(words) <= 1:
            self._c.drawString(x, y, text)
            return
        word_widths = [pdfmetrics.stringWidth(w, font, size) for w in words]
        total_w  = sum(word_widths)
        n_gaps   = len(words) - 1
        gap_w    = (width - total_w) / n_gaps
        cx = x
        for i, (word, ww) in enumerate(zip(words, word_widths)):
            self._c.drawString(cx, y, word)
            if i < n_gaps:
                cx += ww + gap_w

    def body_text(self, text: str, indent: float = 0, size: int = BODY_SIZE) -> None:
        c = self._c
        c.setFont(self._body_reg, size)
        c.setFillColorRGB(0.18, 0.18, 0.18)
        avail_w = CONTENT_W - indent
        # Her paragrafı ayrı işle (explicit \n = paragraf arası)
        for para in text.split("\n"):
            para = para.strip()
            if not para:
                self.y -= LINE_H_BODY // 2
                continue
            lines = wrap_text(para, self._body_reg, size, avail_w)
            for i, ln in enumerate(lines):
                self._ensure(LINE_H_BODY)
                is_last = (i == len(lines) - 1)
                if is_last:
                    c.drawString(MARGIN + indent, self.y, ln)
                else:
                    self._justify_line(ln, MARGIN + indent, self.y, avail_w,
                                       self._body_reg, size)
                self.y -= LINE_H_BODY

    def _get_run_font(self, bold: bool, italic: bool) -> str:
        if bold and italic:
            return self._body_bold_italic
        if bold:
            return self._body_bold
        if italic:
            return self._body_italic
        return self._body_reg

    def body_text_rich(self, paragraphs: list, indent: float = 0) -> None:
        """Yapılandırılmış zengin metin: [{bullet, runs:[{text, bold, italic}]}]"""
        c = self._c
        BULLET_INDENT = 14
        SP_W = pdfmetrics.stringWidth(" ", self._body_reg, BODY_SIZE)

        for para in paragraphs:
            is_bullet = para.get("bullet", False)
            runs      = para.get("runs", [])

            if not runs or all(not r.get("text", "").strip() for r in runs):
                self.y -= LINE_H_BODY // 2
                continue

            x_origin = MARGIN + indent + (BULLET_INDENT if is_bullet else 0)
            eff_w    = CONTENT_W - indent - (BULLET_INDENT if is_bullet else 0)

            # Kelime/boşluk token listesi
            tokens: list = []
            for run in runs:
                fn = self._get_run_font(run.get("bold", False), run.get("italic", False))
                parts = run["text"].split(" ")
                for i, word in enumerate(parts):
                    if i > 0:
                        tokens.append({"t": " ", "fn": fn,
                                       "w": pdfmetrics.stringWidth(" ", fn, BODY_SIZE),
                                       "sp": True})
                    if word:
                        tokens.append({"t": word, "fn": fn,
                                       "w": pdfmetrics.stringWidth(word, fn, BODY_SIZE),
                                       "sp": False})

            # Satırlara yerleştir
            lines: list = []
            cur: list   = []
            cur_w       = 0.0
            for tok in tokens:
                if tok["w"] + cur_w > eff_w and cur and not tok["sp"]:
                    while cur and cur[-1]["sp"]:
                        cur.pop()
                    if cur:
                        lines.append(cur)
                    cur   = [tok]
                    cur_w = tok["w"]
                else:
                    cur.append(tok)
                    cur_w += tok["w"]
            while cur and cur[-1]["sp"]:
                cur.pop()
            if cur:
                lines.append(cur)

            # Çiz
            for li, line_toks in enumerate(lines):
                self._ensure(LINE_H_BODY)
                is_last = (li == len(lines) - 1)
                n_sp    = sum(1 for t in line_toks if t["sp"])
                word_w  = sum(t["w"] for t in line_toks if not t["sp"])

                if not is_last and n_sp > 0:
                    each_sp = (eff_w - word_w) / n_sp
                else:
                    each_sp = SP_W

                if is_bullet and li == 0:
                    c.setFont(self._body_reg, BODY_SIZE)
                    c.setFillColorRGB(0.18, 0.18, 0.18)
                    c.drawString(MARGIN + indent, self.y, "•")

                x = x_origin
                for tok in line_toks:
                    c.setFillColorRGB(0.18, 0.18, 0.18)
                    if tok["sp"]:
                        x += each_sp
                    else:
                        c.setFont(tok["fn"], BODY_SIZE)
                        c.drawString(x, self.y, tok["t"])
                        x += tok["w"]

                self.y -= LINE_H_BODY

    def spacer(self, h: float = 6) -> None:
        self._ensure(h)
        self.y -= h

    def draw_cost_table(self, disciplines: List[Dict]) -> None:
        if not disciplines:
            return
        self._ensure(80)
        self.section_title("Maliyet Tablosu")

        c = self._c
        # Sütun genişlikleri: Disiplin | KDV Hariç | KDV Dahil (%20)
        col_w = [
            CONTENT_W * 0.44,
            CONTENT_W * 0.28,
            CONTENT_W * 0.28,
        ]
        headers = ["Disiplin / Kalem", "KDV Hariç Tutar", "KDV Dahil Tutar (%20)"]
        row_h = 18

        def _draw_row_bg(row_idx, alt=False, h=None):
            rh = h if h is not None else row_h
            if alt:
                bg = hex_rgb("#F8FBFF") if row_idx % 2 == 0 else (1, 1, 1)
            else:
                bg = hex_rgb("#F0F6FF") if row_idx % 2 == 0 else (1, 1, 1)
            self._rect_fill(MARGIN, self.y - rh + 2, CONTENT_W, rh, bg)
            c.setStrokeColorRGB(0.80, 0.85, 0.92)
            c.setLineWidth(0.3)
            c.rect(MARGIN, self.y - rh + 2, CONTENT_W, rh, fill=0, stroke=1)

        # Başlık satırı
        self._ensure(row_h + 4)
        x = MARGIN
        self._rect_fill(MARGIN, self.y - row_h + 2, CONTENT_W, row_h, hex_rgb("#1A2E4A"))
        c.setFont(self._body_bold, 9)
        c.setFillColorRGB(1, 1, 1)
        for i, hdr in enumerate(headers):
            c.drawString(x + 4, self.y - 12, hdr)
            x += col_w[i]
        self.y -= row_h

        total_hariç_val = 0.0
        total_dahil_val = 0.0
        row_idx = 0

        for disc in disciplines:
            alt_kalemler = disc.get("alt_kalemler", [])
            hariç   = disc.get("maliyet_hariç", "")
            dahil   = disc.get("maliyet_dahil", "")
            hariç_v = disc.get("maliyet_hariç_val", 0.0)
            dahil_v = disc.get("maliyet_dahil_val", 0.0)
            total_hariç_val += hariç_v
            total_dahil_val += dahil_v

            if alt_kalemler:
                # Disiplin başlık satırı (tutar yok, sadece isim)
                self._ensure(row_h + 2)
                _draw_row_bg(row_idx)
                c.setFont(self._body_bold, BODY_SIZE - 1)
                c.setFillColorRGB(*hex_rgb("#1A2E4A"))
                c.drawString(MARGIN + 4, self.y - 12, disc.get("name", ""))
                self.y -= row_h
                row_idx += 1

                # Alt kalem satırları (girintili, dinamik yükseklik)
                for ak in alt_kalemler:
                    ak_label  = "↳  " + ak.get("ad", "")
                    ak_lines  = wrap_text(ak_label, self._body_reg,
                                          BODY_SIZE - 1, col_w[0] - 18)
                    n_ak      = max(1, len(ak_lines))
                    ak_row_h  = max(row_h, n_ak * 13 + 8)

                    self._ensure(ak_row_h + 2)
                    _draw_row_bg(row_idx, alt=True, h=ak_row_h)

                    c.setFont(self._body_reg, BODY_SIZE - 1)
                    c.setFillColorRGB(0.18, 0.18, 0.18)
                    # Dikey ortalama
                    text_h      = n_ak * 13
                    first_base  = self.y - (ak_row_h - text_h) / 2 - 11
                    for li, ln in enumerate(ak_lines):
                        c.drawString(MARGIN + 14, first_base - li * 13, ln)

                    mid_y = self.y - ak_row_h / 2 - 4
                    x = MARGIN + col_w[0]
                    c.drawString(x + 4, mid_y, ak.get("maliyet_hariç", ""))
                    x += col_w[1]
                    c.drawString(x + 4, mid_y, ak.get("maliyet_dahil", ""))

                    self.y -= ak_row_h
                    row_idx += 1
            else:
                # Normal disiplin satırı
                self._ensure(row_h + 2)
                _draw_row_bg(row_idx)
                x = MARGIN
                for i, val in enumerate([disc.get("name", ""), hariç, dahil]):
                    c.setFont(self._body_bold if i == 0 else self._body_reg, BODY_SIZE - 1)
                    c.setFillColorRGB(0.1, 0.1, 0.1)
                    c.drawString(x + 4, self.y - 12, val)
                    x += col_w[i]
                self.y -= row_h
                row_idx += 1

        # Toplam satırı (birden fazla disiplin varsa)
        if len(disciplines) > 1 and total_hariç_val:
            self._ensure(row_h + 2)
            self._rect_fill(MARGIN, self.y - row_h + 2, CONTENT_W, row_h, hex_rgb("#E8F0FE"))
            c.setStrokeColorRGB(0.60, 0.70, 0.85)
            c.setLineWidth(0.6)
            c.rect(MARGIN, self.y - row_h + 2, CONTENT_W, row_h, fill=0, stroke=1)
            c.setFont(self._body_bold, BODY_SIZE - 1)
            c.setFillColorRGB(*hex_rgb("#1A2E4A"))
            c.drawString(MARGIN + 4, self.y - 12, "TOPLAM")
            x = MARGIN + col_w[0]
            c.drawString(x + 4, self.y - 12, _fmt_tl(total_hariç_val))
            x += col_w[1]
            c.drawString(x + 4, self.y - 12, _fmt_tl(total_dahil_val))
            self.y -= row_h

        self.spacer(8)

    def draw_photos(self, photos: List[Path],
                    grid=None, crop: bool = True) -> None:
        if not photos:
            return
        cols = grid[0] if grid else 2
        gap  = 6
        cell_w = (CONTENT_W - (cols - 1) * gap) / cols
        cell_h = cell_w * 0.70

        # Section title + ilk fotoğraf satırı aynı sayfada kalmalı
        title_block = 20 + 14 + 4   # box_h + gap_after + spacer
        self.spacer(8)
        self._ensure(title_block + cell_h + gap)
        self.section_title("Fotoğraflar")
        self.spacer(4)
        for idx, ph in enumerate(photos):
            col_idx = idx % cols
            if col_idx == 0:
                self._ensure(cell_h + gap + 10)
            x = MARGIN + col_idx * (cell_w + gap)
            y = self.y - cell_h
            try:
                with Image.open(ph) as im:
                    im = ImageOps.exif_transpose(im)
                    im = im.convert("RGB")
                    iw, ih = im.size
                    if crop:
                        scale = max(cell_w / iw, cell_h / ih)
                        nw = max(1, int(iw * scale))
                        nh = max(1, int(ih * scale))
                        im = im.resize((nw, nh), Image.LANCZOS)
                        lft = (nw - int(cell_w)) // 2
                        top = (nh - int(cell_h)) // 2
                        im = im.crop((lft, top, lft + int(cell_w), top + int(cell_h)))
                        nw, nh = int(cell_w), int(cell_h)
                        ox, oy = 0, 0
                    else:
                        scale = min(cell_w / iw, cell_h / ih)
                        nw = max(1, int(iw * scale))
                        nh = max(1, int(ih * scale))
                        im = im.resize((nw, nh), Image.LANCZOS)
                        ox = (cell_w - nw) / 2
                        oy = (cell_h - nh) / 2
                    buf = io.BytesIO()
                    im.save(buf, "JPEG", quality=88)
                    buf.seek(0)
                    reader = ImageReader(buf)
                    cv = self._c
                    cv.setFillColorRGB(0.95, 0.95, 0.95)
                    cv.rect(x, y, cell_w, cell_h, fill=1, stroke=0)
                    cv.drawImage(reader, x + ox, y + oy, width=nw, height=nh,
                                 preserveAspectRatio=False)
                    cv.setStrokeColorRGB(0.75, 0.80, 0.88)
                    cv.setLineWidth(0.5)
                    cv.rect(x, y, cell_w, cell_h, fill=0, stroke=1)
            except Exception:
                pass
            if col_idx == cols - 1:
                self.y -= cell_h + gap
        if len(photos) % cols != 0:
            self.y -= cell_h + gap

    def draw_footer_note(self, note: str) -> None:
        self.spacer(10)
        c = self._c
        self._rect_fill(MARGIN, self.y - 26, CONTENT_W, 28, hex_rgb("#FFF7ED"))
        self._hline(self.y + 2, 0.8, hex_rgb("#D97706"))
        c.setFont(self._body_reg, 9)
        c.setFillColorRGB(*hex_rgb("#92400E"))
        lines = wrap_text(note, self._body_reg, 9, CONTENT_W - 8)
        for i, ln in enumerate(lines):
            c.drawString(MARGIN + 4, self.y - 10 - i * 11, ln)
        self.y -= 34

    def draw_date_signature(self, tarih: str, duzenleyen: str = "") -> None:
        self.spacer(18)
        need = LINE_H_BODY * (2 if duzenleyen else 1) + 4
        self._ensure(need)
        c = self._c
        c.setFont(self._body_reg, BODY_SIZE)
        c.setFillColorRGB(0.3, 0.3, 0.3)
        # Düzenleyen üstte, tarih altında — ikisi de sağa hizalı
        if duzenleyen:
            duz = f"Düzenleyen: {duzenleyen}"
            dw = pdfmetrics.stringWidth(duz, self._body_reg, BODY_SIZE)
            c.drawString(PAGE_W - MARGIN - dw, self.y, duz)
            self.y -= LINE_H_BODY
        tarih_txt = f"Tarih: {tarih}"
        tw = pdfmetrics.stringWidth(tarih_txt, self._body_reg, BODY_SIZE)
        c.drawString(PAGE_W - MARGIN - tw, self.y, tarih_txt)
        self.y -= LINE_H_BODY

    def finalize(self) -> None:
        self._c.save()


def _fmt_tl(value: float) -> str:
    """float → Türkçe formatında TL string. 150000.0 → '150.000,00 TL'"""
    if value == 0:
        return ""
    s = f"{value:,.2f}"          # '150,000.00'
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return s + " TL"


def _add_page_numbers(pdf_buf: io.BytesIO) -> io.BytesIO:
    """Her sayfaya 'Sayfa X / Y' ekler (post-process, pypdf merge)."""
    pdf_buf.seek(0)
    reader   = PdfReader(pdf_buf)
    n_pages  = len(reader.pages)
    body_reg, _ = get_fonts()

    overlay_buf = io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf, pagesize=A4)
    for i in range(n_pages):
        txt = f"Sayfa {i + 1} / {n_pages}"
        c.setFont(body_reg, 8)
        c.setFillColorRGB(0.55, 0.55, 0.55)
        tw = pdfmetrics.stringWidth(txt, body_reg, 8)
        c.drawString(PAGE_W - MARGIN - tw, MARGIN * 0.45, txt)
        c.showPage()
    c.save()

    overlay_buf.seek(0)
    overlay = PdfReader(overlay_buf)
    writer  = PdfWriter()
    for i, page in enumerate(reader.pages):
        if i < len(overlay.pages):
            page.merge_page(overlay.pages[i])
        writer.add_page(page)

    merged = io.BytesIO()
    writer.write(merged)
    merged.seek(0)
    return merged


def _append_external_pdfs(main_buf: io.BytesIO,
                           ext_paths: List[str]) -> io.BytesIO:
    valids = [p for p in ext_paths if p and Path(p).exists()]
    if not valids:
        return main_buf
    main_buf.seek(0)
    writer = PdfWriter()
    writer.append(main_buf)
    for ep in valids:
        writer.append(ep)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out


# ──────────────────────────────────────────────────────────────────────────────
# Ana üretim fonksiyonu
# ──────────────────────────────────────────────────────────────────────────────

def build_note(data: Dict[str, Any]) -> Dict[str, str]:
    konu             = data.get("konu", "Bilgi Notu")
    aciklamalar_rich = data.get("aciklamalar_rich", [])
    aciklamalar      = data.get("aciklamalar", "")   # düz metin geri uyumluluk
    talep_eden       = data.get("talep_eden")
    disciplines = data.get("disciplines", [])
    images      = [Path(p) for p in data.get("images", [])]
    photo_grid  = data.get("photo_grid")
    photo_crop  = data.get("photo_crop", True)
    ext_pdfs    = data.get("ext_pdfs", [])
    tarih       = data.get("tarih") or _today_str()
    duzenleyen  = data.get("duzenleyen", "")
    fname_base  = _safe_fname(konu)

    out_dir = Path.home() / "Desktop"
    out_dir.mkdir(parents=True, exist_ok=True)

    # ── PDF ───────────────────────────────────────────────────────────────────
    pdf_buf = io.BytesIO()
    p = _PDFPainter(pdf_buf)

    p.draw_main_header()
    p.spacer(10)

    p.section_title("Konu")
    p.body_text(konu, indent=4)
    p.spacer(8)

    if talep_eden:
        p.section_title("Talep Eden Birim")
        p.body_text(talep_eden, indent=4)
        p.spacer(8)

    if aciklamalar_rich:
        p.section_title("Açıklamalar")
        p.body_text_rich(aciklamalar_rich, indent=4)
        p.spacer(10)
    elif aciklamalar:
        p.section_title("Açıklamalar")
        p.body_text(aciklamalar, indent=4)
        p.spacer(10)

    if disciplines:
        p.draw_cost_table(disciplines)

    p.draw_date_signature(tarih, duzenleyen)

    p.draw_photos(images, grid=photo_grid, crop=photo_crop)

    if _need_saha_note(disciplines):
        p.draw_footer_note(_SAHA_NOTE)

    p.finalize()

    # Sayfa numaraları (Sayfa X / Y) — tüm sayfalar bilinince eklenir
    pdf_buf.seek(0)
    numbered_buf = _add_page_numbers(pdf_buf)
    final_buf = _append_external_pdfs(numbered_buf, ext_pdfs)

    pdf_path = str(out_dir / (fname_base + ".pdf"))
    with open(pdf_path, "wb") as fh:
        fh.write(final_buf.read())

    docx_path = _build_docx(data, tarih, duzenleyen, out_dir, fname_base)
    return {"pdf": pdf_path, "docx": docx_path}


# ──────────────────────────────────────────────────────────────────────────────
# Word üretimi
# ──────────────────────────────────────────────────────────────────────────────

def _build_docx(data: Dict, tarih: str, duzenleyen: str,
                out_dir: Path, fname_base: str) -> str:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    konu             = data.get("konu", "Bilgi Notu")
    aciklamalar_rich = data.get("aciklamalar_rich", [])
    aciklamalar      = data.get("aciklamalar", "")
    talep_eden       = data.get("talep_eden")
    disciplines      = data.get("disciplines", [])
    images           = [Path(p) for p in data.get("images", [])]

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    TNR = "Times New Roman"

    def _set_run_font(run, size: float = 12, bold: bool = False,
                      color: RGBColor = None, italic: bool = False):
        run.font.name = TNR
        run.font.size = Pt(size)
        run.bold      = bold
        run.italic    = italic
        if color:
            run.font.color.rgb = color

    # ── Logo + Başlık ─────────────────────────────────────────────────────────
    logo = _logo_path()
    hdr_para = doc.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo:
        try:
            run = hdr_para.add_run()
            run.add_picture(str(logo), height=Cm(1.4))
            hdr_para.add_run("  ")
        except Exception:
            pass

    run_birim = hdr_para.add_run(BIRIM_ADI.replace("\n", " · "))
    _set_run_font(run_birim, size=9, color=RGBColor(0x47, 0x49, 0x4F))
    hdr_para.add_run("    ")
    run_belge = hdr_para.add_run(BELGE_ADI)
    _set_run_font(run_belge, size=14, bold=True, color=RGBColor(0x1A, 0x2E, 0x4A))

    # Mavi yatay çizgi
    def _add_blue_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2563EB")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(4)

    _add_blue_hr()

    def _section_heading(text: str):
        """Bölüm başlığı: lacivert kalın TNR + paragraf alt çizgisi."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)   # title case olarak geçirilir, upper() yok
        _set_run_font(run, size=12, bold=True, color=RGBColor(0x1A, 0x2E, 0x4A))
        # Alt ince lacivert çizgi
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "4A90D9")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _body(text: str, bold: bool = False, size: float = 12,
              color: RGBColor = None, italic: bool = False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, color=color, italic=italic)
        p.paragraph_format.space_after  = Pt(2)
        p.alignment = align
        return p

    # ── Konu ──────────────────────────────────────────────────────────────────
    _section_heading("Konu")
    _body(konu)

    # ── Talep Eden Birim ──────────────────────────────────────────────────────
    if talep_eden:
        _section_heading("Talep Eden Birim")
        _body(talep_eden)

    # ── Açıklamalar ───────────────────────────────────────────────────────────
    if aciklamalar_rich:
        _section_heading("Açıklamalar")
        for para in aciklamalar_rich:
            runs_data = [r for r in para.get("runs", []) if r.get("text")]
            if not runs_data and not para.get("bullet"):
                # Boş paragraf → küçük boşluk olarak ekle
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)
                continue
            p_doc = doc.add_paragraph()
            p_doc.paragraph_format.space_after  = Pt(3)
            p_doc.paragraph_format.space_before = Pt(0)
            p_doc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if para.get("bullet"):
                p_doc.paragraph_format.left_indent    = Cm(0.5)
                p_doc.paragraph_format.first_line_indent = Cm(-0.5)
                bullet_run = p_doc.add_run("• ")
                _set_run_font(bullet_run, size=11)
            for rd in runs_data:
                r = p_doc.add_run(rd["text"])
                _set_run_font(r, size=11,
                              bold=rd.get("bold", False),
                              italic=rd.get("italic", False))
    elif aciklamalar:
        _section_heading("Açıklamalar")
        _body(aciklamalar)

    # ── Maliyet Tablosu ───────────────────────────────────────────────────────
    if disciplines:
        _section_heading("Maliyet Tablosu")
        col_headers = ["Disiplin / Kalem", "KDV Hariç Tutar", "KDV Dahil Tutar (%20)"]

        all_rows = []
        for disc in disciplines:
            alt_kalemler = disc.get("alt_kalemler", [])
            if alt_kalemler:
                all_rows.append(("disc_head", disc))
                for ak in alt_kalemler:
                    all_rows.append(("sub", ak))
            else:
                all_rows.append(("disc", disc))
        total_h = sum(d.get("maliyet_hariç_val", 0.0) for d in disciplines)
        total_d = sum(d.get("maliyet_dahil_val", 0.0) for d in disciplines)
        if len(disciplines) > 1 and total_h:
            all_rows.append(("total", total_h, total_d))

        tbl = doc.add_table(rows=1 + len(all_rows), cols=3)
        tbl.style = "Table Grid"

        from docx.oxml import OxmlElement as _OE
        from docx.shared import Cm as _Cm

        # Sütun genişlikleri: sayfanın kullanılabilir genişliği ~16cm
        col_widths = [_Cm(7.5), _Cm(4.25), _Cm(4.25)]
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

        def _dark_cell(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = _OE("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1A2E4A")
            tcPr.append(shd)

        def _light_cell(cell, fill="E8F0FE"):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = _OE("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

        # Başlık satırı
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(col_headers):
            hdr_cells[i].text = ""
            run = hdr_cells[i].paragraphs[0].add_run(h)
            _set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            _dark_cell(hdr_cells[i])

        for row_i, row_data in enumerate(all_rows):
            cells = tbl.rows[row_i + 1].cells
            rtype = row_data[0]

            if rtype == "disc":
                disc = row_data[1]
                cells[0].text = ""
                run = cells[0].paragraphs[0].add_run(disc.get("name", ""))
                _set_run_font(run, size=11, bold=True)
                cells[1].text = ""
                run = cells[1].paragraphs[0].add_run(disc.get("maliyet_hariç", ""))
                _set_run_font(run, size=11)
                cells[2].text = ""
                run = cells[2].paragraphs[0].add_run(disc.get("maliyet_dahil", ""))
                _set_run_font(run, size=11)

            elif rtype == "disc_head":
                disc = row_data[1]
                cells[0].text = ""
                run = cells[0].paragraphs[0].add_run(disc.get("name", ""))
                _set_run_font(run, size=11, bold=True)
                for j in [1, 2]:
                    cells[j].text = ""

            elif rtype == "sub":
                ak = row_data[1]
                cells[0].text = ""
                run = cells[0].paragraphs[0].add_run("    ↳  " + ak.get("ad", ""))
                _set_run_font(run, size=11, color=RGBColor(0x33, 0x33, 0x33))
                cells[1].text = ""
                run = cells[1].paragraphs[0].add_run(ak.get("maliyet_hariç", ""))
                _set_run_font(run, size=11)
                cells[2].text = ""
                run = cells[2].paragraphs[0].add_run(ak.get("maliyet_dahil", ""))
                _set_run_font(run, size=11)

            elif rtype == "total":
                _, th, td = row_data
                cells[0].text = ""
                run = cells[0].paragraphs[0].add_run("TOPLAM")
                _set_run_font(run, size=11, bold=True, color=RGBColor(0x1A, 0x2E, 0x4A))
                _light_cell(cells[0])
                cells[1].text = ""
                run = cells[1].paragraphs[0].add_run(_fmt_tl(th))
                _set_run_font(run, size=11, bold=True)
                _light_cell(cells[1])
                cells[2].text = ""
                run = cells[2].paragraphs[0].add_run(_fmt_tl(td))
                _set_run_font(run, size=11, bold=True)
                _light_cell(cells[2])

    # ── Tarih + Düzenleyen (fotoğraflardan önce) ─────────────────────────────
    doc.add_paragraph()
    clr = RGBColor(0x33, 0x33, 0x33)
    if duzenleyen:
        p_duz = doc.add_paragraph()
        p_duz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_duz = p_duz.add_run(f"Düzenleyen: {duzenleyen}")
        _set_run_font(run_duz, size=12, color=clr)
        p_duz.paragraph_format.space_after = Pt(0)
    p_tar = doc.add_paragraph()
    p_tar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_tar = p_tar.add_run(f"Tarih: {tarih}")
    _set_run_font(run_tar, size=12, color=clr)

    # ── Fotoğraflar ───────────────────────────────────────────────────────────
    if images:
        _section_heading("Fotoğraflar")
        tbl = doc.add_table(rows=0, cols=2)
        for i in range(0, len(images), 2):
            row = tbl.add_row().cells
            for j in range(2):
                if i + j < len(images):
                    ph = images[i + j]
                    try:
                        para = row[j].paragraphs[0]
                        run = para.add_run()
                        run.add_picture(str(ph), width=Cm(8.0))
                    except Exception:
                        row[j].text = ph.name

    # ── Saha notu ─────────────────────────────────────────────────────────────
    if _need_saha_note(disciplines):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(_SAHA_NOTE)
        _set_run_font(run, size=10, italic=True,
                      color=RGBColor(0x92, 0x40, 0x0E))

    docx_path = str(out_dir / (fname_base + ".docx"))
    doc.save(docx_path)
    return docx_path
